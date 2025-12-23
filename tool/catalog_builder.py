
from __future__ import annotations

import os
import re
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
import openpyxl
from docx import Document

# ====== Helpers ======
def _clean(x) -> str:
    if x is None:
        return ""
    try:
        if isinstance(x, float) and np.isnan(x):
            return ""
    except Exception:
        pass
    return str(x).strip()

def _norm_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", _clean(s)).strip()

def _semester_from_value(v) -> str:
    s = _clean(v).lower()
    if not s:
        return ""
    if s in {"1","hk1","hki","i"}:
        return "HK1"
    if s in {"2","hk2","hkii","ii"}:
        return "HK2"
    if "hk1" in s or "học kì i" in s or "hoc ki i" in s:
        return "HK1"
    if "hk2" in s or "học kì ii" in s or "hoc ki ii" in s:
        return "HK2"
    if re.fullmatch(r"\d+", s):
        return "HK1" if s == "1" else ("HK2" if s == "2" else "")
    return ""

def _parse_grade_from_filename(name: str) -> Optional[int]:
    m = re.search(r"([1-5])", name)
    return int(m.group(1)) if m else None

def _find_header_rows(df: pd.DataFrame) -> List[int]:
    hdr: List[int] = []
    for i in range(len(df)):
        row = [_clean(x) for x in df.iloc[i].tolist()]
        joined = " ".join(row).lower()
        if ("yêu cầu" in joined or "ycc" in joined) and ("bài" in joined or "các bài" in joined):
            if any(k in joined for k in ["chủ đề", "tên chủ đề", "chủ điểm", "phần", "mạch"]):
                hdr.append(i)
    if not hdr:
        for i in range(len(df)):
            joined = " ".join([_clean(x) for x in df.iloc[i].tolist()]).lower()
            if ("yêu cầu" in joined) and ("bài" in joined):
                hdr.append(i)
    # de-dup
    out: List[int] = []
    for x in hdr:
        if x not in out:
            out.append(x)
    return out

def _map_cols(header_row: List[str]) -> Dict[str, int]:
    cols: Dict[str, int] = {}
    for idx, val in enumerate(header_row):
        v = _clean(val).lower()
        if not v:
            continue
        if "học kì" in v or v in {"hk","hoc ki","semester"}:
            cols["semester"] = idx
        elif any(k in v for k in ["tên chủ đề","chủ đề","chủ điểm","phần"]):
            cols.setdefault("topic", idx)
        elif v.strip() == "bài":
            cols["bai"] = idx
        elif any(k in v for k in ["tên bài","các bài","bài học cụ thể","nội dung"]):
            cols["lesson"] = idx
        elif "yêu cầu" in v or "ycc" in v:
            cols["yccd"] = idx
        elif v.strip() in {"stt","tt","số tt","so tt"}:
            cols["stt"] = idx
    return cols

def _infer_semester_grade2_toan(topic: str) -> str:
    # "Chủ đề 7: Ôn tập học kì 1" -> HK1, topics >7 -> HK2
    t = _clean(topic)
    m = re.search(r"Chủ\s*đề\s*(\d+)", t, flags=re.I)
    if not m:
        return ""
    num = int(m.group(1))
    return "HK1" if num <= 7 else "HK2"

def _semester_tv2(topic: str, order: List[str]) -> str:
    if not order:
        return ""
    half = (len(order) + 1) // 2
    try:
        idx = order.index(topic)
    except ValueError:
        return ""
    return "HK1" if idx < half else "HK2"

def _subject_from_sheet(sheet: str) -> Optional[str]:
    s = _clean(sheet)
    mapping = {
        "TV": "Tiếng Việt",
        "TV2": "Tiếng Việt",
        "TOán": "Toán",
        "Toán": "Toán",
        "Toán 2": "Toán",
        "Khoa học": "Khoa học",
        "LSĐL": "Lịch sử - Địa lý",
        "Công nghệ": "Công nghệ",
        "Tin học": "Tin",
        "Tin học CKP": "Tin",
    }
    return mapping.get(s)

def parse_xlsx_to_catalog(xlsx_path: str) -> pd.DataFrame:
    """Parse the provided KHGD excel format into a normalized YCCĐ catalog."""
    grade = _parse_grade_from_filename(os.path.basename(xlsx_path))
    if grade is None:
        raise ValueError(f"Không xác định được lớp từ tên file: {os.path.basename(xlsx_path)}")

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    out_rows: List[dict] = []

    for sheet in wb.sheetnames:
        subject = _subject_from_sheet(sheet)
        if not subject:
            continue
        # pandas read
        df = pd.read_excel(xlsx_path, sheet_name=sheet, header=None)
        if df is None or df.empty:
            continue

        hdrs = _find_header_rows(df)
        if not hdrs:
            continue

        # TV2 topic order (for semester splitting)
        tv2_order: List[str] = []
        if grade == 2 and subject == "Tiếng Việt":
            # header row 0; topic in col 1 usually
            for i in range(hdrs[0] + 1, len(df)):
                t = _clean(df.iloc[i, 1] if df.shape[1] > 1 else "")
                if t and t.lower() not in {"tên chủ đề"} and t not in tv2_order:
                    tv2_order.append(t)

        for seg_idx, h in enumerate(hdrs):
            header = [_clean(x) for x in df.iloc[h].tolist()]
            cols = _map_cols(header)
            start = h + 1
            end = (hdrs[seg_idx + 1] - 1) if seg_idx + 1 < len(hdrs) else (len(df) - 1)

            seg_sem = ""
            if "semester" not in cols and len(hdrs) >= 2:
                if seg_idx == 0:
                    seg_sem = "HK1"
                elif seg_idx == 1:
                    seg_sem = "HK2"

            last_topic = ""
            last_lesson = ""
            last_bai = ""
            for r in range(start, end + 1):
                row = df.iloc[r].tolist()
                joined = " ".join([_clean(x) for x in row]).lower()
                # skip accidental duplicated header rows
                if ("yêu cầu" in joined and "bài" in joined and ("chủ đề" in joined or "chủ điểm" in joined)) and r != start:
                    continue

                topic = _clean(row[cols["topic"]]) if "topic" in cols else ""
                if topic:
                    last_topic = topic
                else:
                    topic = last_topic

                lesson = _clean(row[cols["lesson"]]) if "lesson" in cols else ""
                if lesson:
                    last_lesson = lesson
                else:
                    lesson = last_lesson

                bai = _clean(row[cols["bai"]]) if "bai" in cols else ""
                if bai:
                    last_bai = bai
                else:
                    bai = last_bai

                yccd = _clean(row[cols["yccd"]]) if "yccd" in cols else ""
                if not yccd:
                    continue

                semester = ""
                if "semester" in cols:
                    semester = _semester_from_value(row[cols["semester"]])
                if not semester:
                    if grade == 2 and subject == "Toán":
                        semester = _infer_semester_grade2_toan(topic)
                    elif grade == 2 and subject == "Tiếng Việt":
                        semester = _semester_tv2(topic, tv2_order)
                    else:
                        semester = seg_sem

                lesson_str = _norm_spaces(lesson)
                if bai and bai != "-" and lesson_str and not re.match(r"^\s*Bài", lesson_str, flags=re.I):
                    lesson_str = _norm_spaces(f"Bài {bai}: {lesson_str}")
                elif bai and bai != "-" and not lesson_str:
                    lesson_str = f"Bài {bai}"

                out_rows.append({
                    "grade": grade,
                    "subject": subject,
                    "semester": semester,
                    "topic": _norm_spaces(topic),
                    "lesson": _norm_spaces(lesson_str),
                    "yccd": _norm_spaces(yccd),
                })

    df_out = pd.DataFrame(out_rows)
    if df_out.empty:
        return df_out

    # normalize "Chủ đề A." -> "Chủ đề A:"
    df_out["topic"] = df_out["topic"].astype(str).str.replace(r"Chủ\s*đề\s*([A-F])\.", r"Chủ đề \1:", regex=True)
    df_out = df_out.drop_duplicates(subset=["grade","subject","semester","topic","lesson","yccd"]).reset_index(drop=True)
    return df_out

def parse_tin_doc_grade5(doc_path: str) -> pd.DataFrame:
    """Fallback Tin lớp 5 from docx (no semester info -> blank semester)."""
    doc = Document(doc_path)
    if len(doc.tables) < 4:
        return pd.DataFrame(columns=["grade","subject","semester","topic","lesson","yccd"])
    t = doc.tables[3]  # grade 5 details in provided doc
    rows = []
    current_topic = ""
    for ridx, row in enumerate(t.rows):
        if ridx == 0:
            continue
        c0 = _norm_spaces(row.cells[0].text)
        c1 = _norm_spaces(row.cells[1].text if len(row.cells) > 1 else "")
        if not c0:
            continue
        if c0.lower().startswith("chủ đề"):
            current_topic = c0.replace("Chủ đề A.", "Chủ đề A:").replace("Chủ đề B.", "Chủ đề B:").replace("Chủ đề C.", "Chủ đề C:") \
                             .replace("Chủ đề D.", "Chủ đề D:").replace("Chủ đề E.", "Chủ đề E:").replace("Chủ đề F.", "Chủ đề F:")
            continue
        if not c1:
            continue
        rows.append({
            "grade": 5,
            "subject": "Tin",
            "semester": "",
            "topic": current_topic,
            "lesson": c0,
            "yccd": c1
        })
    return pd.DataFrame(rows)

def build_catalog_from_sources(source_dir: str, output_csv: Optional[str] = None) -> pd.DataFrame:
    """Build catalog from KHGD sources (xlsx)."""
    xlsx_paths = [os.path.join(source_dir, f) for f in os.listdir(source_dir) if f.lower().endswith(".xlsx")]
    frames: List[pd.DataFrame] = []
    for p in sorted(xlsx_paths):
        try:
            frames.append(parse_xlsx_to_catalog(p))
        except Exception:
            # ignore broken files
            continue
    df = pd.concat(frames, ignore_index=True) if frames else pd.DataFrame(columns=["grade","subject","semester","topic","lesson","yccd"])

    # Tin lớp 5 fallback if present
    tin_doc = os.path.join(source_dir, "04. Yeu cau can dat mon TIN HOC.docx")
    if os.path.exists(tin_doc):
        df2 = parse_tin_doc_grade5(tin_doc)
        if not df2.empty:
            df = pd.concat([df, df2], ignore_index=True)

    # final cleanup
    for c in ["subject","semester","topic","lesson","yccd"]:
        if c not in df.columns:
            df[c] = ""
        df[c] = df[c].fillna("").astype(str).map(_norm_spaces)
    df["grade"] = pd.to_numeric(df.get("grade"), errors="coerce").astype("Int64")
    df = df.dropna(subset=["grade"]).drop_duplicates(subset=["grade","subject","semester","topic","lesson","yccd"]).reset_index(drop=True)

    if output_csv:
        os.makedirs(os.path.dirname(output_csv) or ".", exist_ok=True)
        df.to_csv(output_csv, index=False, encoding="utf-8-sig")
    return df

def load_or_build_catalog(csv_path: str, source_dir: str) -> pd.DataFrame:
    if os.path.exists(csv_path):
        try:
            return pd.read_csv(csv_path)
        except Exception:
            pass
    return build_catalog_from_sources(source_dir, output_csv=csv_path)
