\
from __future__ import annotations
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional
from copy import deepcopy
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .utils import QTYPE_ORDER, LEVEL_ORDER, fmt_ranges
from .generation import Slot, slot_map_to_numbers, totals_by_cell
from .matrix_template import MatrixTemplate

DOCX_QTYPE_TO_COL_START = {
    "MCQ": 4,    # columns 4-6
    "TF": 7,     # 7-9
    "MATCH": 10, # 10-12
    "FILL": 13,  # 13-15
    "ESSAY": 16, # 16-18
}

def _delete_row(table, row_idx: int):
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    tbl.remove(tr)

def _clone_row(table, row_idx: int):
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    table._tbl.insert(row_idx+1, new_tr)

def export_spec_from_template(
    template_docx_path: str,
    output_path: str,
    matrix: MatrixTemplate,
    slots: List[Slot],
    points_per_qtype: Dict[str, float],
    total_points: float = 10.0
) -> str:
    doc = Document(template_docx_path)
    table = doc.tables[0]

    # Identify header rows (0-3), totals rows last 3
    header_rows = 4
    totals_rows = 3
    current_rows = len(table.rows)
    data_start = header_rows
    data_end = current_rows - totals_rows  # exclusive

    # Delete existing data rows
    # delete from bottom to top to keep indices stable
    for ridx in range(data_end-1, data_start-1, -1):
        _delete_row(table, ridx)

    # Use the first removed template row as style base by cloning current row at data_start if exists,
    # else clone last header row.
    style_base_idx = data_start if len(table.rows) > data_start else header_rows-1
    # Insert lesson rows by cloning style row each time, then fill.
    # We need exactly len(matrix.lessons) rows before totals rows.
    for i, lesson in enumerate(matrix.lessons):
        # Add a new row at the end (before totals) by cloning the style base row.
        # Easiest: add_row then try to copy style from base cells.
        row = table.add_row()
        # Copy formatting by cloning XML of base row (safer). We'll replace the last row with clone of base row.
        base_tr = table.rows[style_base_idx]._tr
        new_tr = deepcopy(base_tr)
        table._tbl.remove(row._tr)
        table._tbl.insert(len(table.rows)-totals_rows, new_tr)

    # Now fill lesson rows
    numbers_map = slot_map_to_numbers(slots)
    # lesson rows are from data_start to data_start+len(lessons)-1
    for i, lesson in enumerate(matrix.lessons):
        r = data_start + i
        cells = table.rows[r].cells
        # TT / topic / lesson / yccd
        cells[0].text = str(lesson.tt)
        cells[1].text = lesson.topic or ""
        cells[2].text = lesson.lesson or ""
        # YCCĐ: leave blank; user can add from bank later (or keep)
        if len(cells) > 3:
            if cells[3].text.strip() in ("…….", ".......", "……", ""):
                cells[3].text = ""
        # Fill each qtype-level cell with "Câu ..." numbers
        for qtype in QTYPE_ORDER:
            base = DOCX_QTYPE_TO_COL_START[qtype]
            for level in LEVEL_ORDER:
                col = base + (level-1)
                key = (lesson.topic, lesson.lesson, qtype, level)
                nums = numbers_map.get(key, [])
                txt = fmt_ranges(nums)
                if txt:
                    cells[col].text = f"Câu {txt}"
                else:
                    cells[col].text = ""

    # Fill totals rows
    tot_counts = totals_by_cell(matrix)
    # Row indices for totals rows
    r_counts = data_start + len(matrix.lessons)
    r_points = r_counts + 1
    r_ratio = r_counts + 2

    # counts per cell
    for qtype in QTYPE_ORDER:
        base = DOCX_QTYPE_TO_COL_START[qtype]
        for level in LEVEL_ORDER:
            col = base + (level-1)
            c = tot_counts.get((qtype, level), 0)
            table.rows[r_counts].cells[col].text = str(c) if c else ""

    # points per cell
    total_pts = 0.0
    for qtype in QTYPE_ORDER:
        base = DOCX_QTYPE_TO_COL_START[qtype]
        for level in LEVEL_ORDER:
            col = base + (level-1)
            c = tot_counts.get((qtype, level), 0)
            pts = c * float(points_per_qtype.get(qtype, 0.25))
            total_pts += pts
            table.rows[r_points].cells[col].text = f"{pts:.2f}" if pts else ""

    # ratio % per cell (based on points contribution)
    for qtype in QTYPE_ORDER:
        base = DOCX_QTYPE_TO_COL_START[qtype]
        for level in LEVEL_ORDER:
            col = base + (level-1)
            c = tot_counts.get((qtype, level), 0)
            pts = c * float(points_per_qtype.get(qtype, 0.25))
            ratio = (pts / total_points * 100.0) if total_points else 0.0
            table.rows[r_ratio].cells[col].text = f"{ratio:.0f}%" if pts else ""

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def export_exam_docx(
    output_path: str,
    slots: List[Slot],
    bank_df,
    title: str,
    total_points: float = 10.0
) -> str:
    """Create a simple exam docx (no external template)."""
    from docx.shared import Pt
    doc = Document()
    doc.add_paragraph(title).runs[0].bold = True
    doc.add_paragraph(f"Thang điểm: {total_points:g}").italic = True
    doc.add_paragraph("")

    # index bank by id
    idx = {}
    for _, r in bank_df.iterrows():
        idx[str(r["question_id"])] = r

    for s in slots:
        p = doc.add_paragraph()
        p.add_run(f"Câu {s.qno}. ").bold = True
        p.add_run(f"({s.points:g} điểm) ")
        if s.question_id and s.question_id in idx:
            r = idx[s.question_id]
            p.add_run(str(r.get("stem","")))
            qtype = str(r.get("qtype","")).upper()
            if qtype == "MCQ":
                import json
                try:
                    opts = json.loads(r.get("options","[]"))
                except Exception:
                    opts = []
                letters = ["A","B","C","D","E","F"]
                for j, opt in enumerate(opts[:6]):
                    doc.add_paragraph(f"{letters[j]}. {opt}", style=None)
            elif qtype in ("TF","MATCH","FILL"):
                doc.add_paragraph("(GV điền nội dung chi tiết cho dạng câu này trong kho câu hỏi)", style=None)
            else:
                doc.add_paragraph("........................................................................")
        else:
            p.add_run("[CHƯA CÓ CÂU PHÙ HỢP TRONG KHO — vui lòng bổ sung hoặc giảm số câu ô này]")
        doc.add_paragraph("")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
